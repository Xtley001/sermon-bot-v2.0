"""
Scrapes Telegram channels and loads materials into RAG system.
Handles data extraction, validation, and vector store population.
"""
import os
import logging
import asyncio
from datetime import datetime
from typing import List, Dict, Optional
import re
import json

from telethon import TelegramClient
from telethon.tl.types import MessageMediaPhoto
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from docx import Document as DocxDocument
from PyPDF2 import PdfReader

import config
from db_handler import SermonDatabase
from utils import rag_engine, llm

logger = logging.getLogger(__name__)


class ChannelScraper:
    """Scrapes sermon messages from Telegram channels."""
    
    def __init__(self):
        """Initialize Telegram client for scraping."""
        self.client = None
        self.db = SermonDatabase(config.DB_PATH)
    
    async def initialize(self):
        """Connect to Telegram."""
        try:
            self.client = TelegramClient(
                'bot_session',
                config.TELEGRAM_API_ID,
                config.TELEGRAM_API_HASH
            )
            await self.client.start(phone=config.TELEGRAM_PHONE)
            logger.info("Telegram client connected for scraping")
        except Exception as e:
            logger.error(f"Failed to initialize Telegram client: {e}")
            raise
    
    async def scrape_channel(self, channel_username: str, limit: int = 500) -> List[Dict]:
        """
        Scrape messages from a channel.
        Only extracts teaching content (ignores announcements, etc).
        """
        logger.info(f"Scraping channel: {channel_username}")
        sermons = []
        
        try:
            # Get channel entity
            channel = await self.client.get_entity(channel_username)
            
            # Fetch messages
            async for message in self.client.iter_messages(channel, limit=limit):
                # Skip empty messages
                if not message.text:
                    continue
                
                # Check if it's a teaching/sermon (use AI)
                if not await self._is_teaching(message.text):
                    continue
                
                # Extract sermon data
                sermon_data = await self._extract_sermon_data(message, channel_username)
                if sermon_data:
                    sermons.append(sermon_data)
                    
                    # Add to database
                    self.db.add_sermon(sermon_data)
            
            logger.info(f"Scraped {len(sermons)} sermons from {channel_username}")
            return sermons
            
        except Exception as e:
            logger.error(f"Error scraping channel {channel_username}: {e}")
            return []
    
    async def _is_teaching(self, text: str) -> bool:
        """Use AI to determine if message is a teaching/sermon."""
        # Quick filters first
        if len(text) < 100:
            return False
        
        # Keywords that indicate teaching
        teaching_keywords = [
            'message', 'sermon', 'teaching', 'word', 'scripture', 
            'bible', 'god', 'jesus', 'pastor', 'ministry', 'anointing',
            'faith', 'prayer', 'worship', 'spirit', 'church', 'kingdom',
            'testimony', 'revelation', 'prophetic', 'glory', 'grace'
        ]
        
        text_lower = text.lower()
        keyword_count = sum(1 for keyword in teaching_keywords if keyword in text_lower)
        
        # If it has at least 2 teaching keywords, likely a sermon
        if keyword_count < 2:
            return False
        
        # Use AI for final validation (optional - can be disabled for speed)
        try:
            prompt = f"""Is this a sermon/teaching message from a pastor? Answer only YES or NO.

Message: "{text[:400]}..."
"""
            response = llm.invoke(prompt)
            answer = response.content.strip().upper()
            return 'YES' in answer
            
        except Exception as e:
            logger.error(f"Error validating teaching: {e}")
            # If AI fails, use keyword count
            return keyword_count >= 2
    
    async def _extract_sermon_data(self, message, channel_username: str) -> Optional[Dict]:
        """Extract title, description, image, link, date, theme from message."""
        try:
            # Get message link
            message_link = f"https://t.me/{channel_username.replace('@', '')}/{message.id}"
            
            # Check if already exists
            existing = self.db.get_sermon_by_link(message_link)
            if existing:
                logger.debug(f"Sermon already exists: {message_link}")
                return None
            
            # Get image URL if photo
            image_url = None
            if isinstance(message.media, MessageMediaPhoto):
                # For simplicity, use the message link as image reference
                # In production, you could download and host images
                image_url = message_link
            
            # Extract title and description using AI
            text = message.text
            title, description, theme = await self._extract_metadata(text)
            
            # Format date
            date = message.date.strftime("%Y-%m-%d") if message.date else None
            
            return {
                'title': title,
                'description': description,
                'channel': channel_username,
                'message_link': message_link,
                'image_url': image_url,
                'date': date,
                'theme': theme
            }
            
        except Exception as e:
            logger.error(f"Error extracting sermon data: {e}")
            return None
    
    async def _extract_metadata(self, text: str) -> tuple:
        """Use AI to extract title, description, and theme from message text."""
        try:
            # Truncate text if too long for API
            text_sample = text[:2000] if len(text) > 2000 else text
            
            prompt = f"""Extract metadata from this sermon message:

Message: "{text_sample}"

Return a JSON object with:
- "title": A clear, concise title (5-15 words max)
- "description": A full description of the sermon content (30-200 words)
- "theme": Main theme/topic (1-3 words like "Faith", "Healing", "Purpose")

Return ONLY the JSON object, nothing else.

Example:
{{"title": "Walking in Faith During Difficult Times", "description": "Pastor Tara teaches about maintaining faith when facing challenges...", "theme": "Faith"}}
"""
            response = llm.invoke(prompt)
            content = response.content.strip()
            
            # Parse JSON
            # Remove markdown code blocks if present
            if content.startswith('```'):
                content = content.split('```')[1]
                if content.startswith('json'):
                    content = content[4:]
                content = content.strip()
            
            # Remove any trailing backticks
            content = content.rstrip('`').strip()
            
            metadata = json.loads(content)
            
            return (
                metadata.get('title', 'Untitled Sermon')[:200],
                metadata.get('description', text[:500])[:1000],
                metadata.get('theme', 'General')[:50]
            )
            
        except json.JSONDecodeError as e:
            logger.error(f"JSON parsing error in metadata extraction: {e}")
            logger.debug(f"AI response was: {content[:200]}")
            # Fallback: use simple extraction
            return self._fallback_extract_metadata(text)
        except Exception as e:
            logger.error(f"Error extracting metadata with AI: {e}")
            return self._fallback_extract_metadata(text)
    
    def _fallback_extract_metadata(self, text: str) -> tuple:
        """Fallback method to extract metadata without AI."""
        lines = text.split('\n')
        lines = [line.strip() for line in lines if line.strip()]
        
        # Title: Use first non-empty line or first sentence
        title = "Untitled Sermon"
        if lines:
            title = lines[0][:200]
            # If first line is too short, try first sentence
            if len(title) < 20 and len(text) > 50:
                sentences = text.split('.')
                if sentences:
                    title = sentences[0][:200]
        
        # Description: Use first 500 chars
        description = text[:500].replace('\n', ' ').strip()
        
        # Theme: Extract from common patterns
        theme = "General"
        theme_keywords = {
            'faith': 'Faith',
            'healing': 'Healing',
            'prosperity': 'Prosperity',
            'purpose': 'Purpose',
            'prayer': 'Prayer',
            'worship': 'Worship',
            'family': 'Family',
            'business': 'Business',
            'breakthrough': 'Breakthrough',
            'deliverance': 'Deliverance',
            'grace': 'Grace',
            'love': 'Love',
            'power': 'Power',
            'supernatural': 'Supernatural'
        }
        
        text_lower = text.lower()
        for keyword, theme_name in theme_keywords.items():
            if keyword in text_lower:
                theme = theme_name
                break
        
        return (title, description, theme)
    
    async def scrape_all_channels(self):
        """Scrape all configured channels."""
        await self.initialize()
        
        all_sermons = []
        for channel in config.CHANNELS_TO_SCRAPE:
            try:
                sermons = await self.scrape_channel(channel)
                all_sermons.extend(sermons)
            except Exception as e:
                logger.error(f"Failed to scrape {channel}: {e}")
                continue
        
        # Add to vector store
        if all_sermons:
            documents = self._sermons_to_documents(all_sermons)
            rag_engine.add_documents(documents)
        
        await self.client.disconnect()
        
        logger.info(f"Total sermons scraped: {len(all_sermons)}")
        return all_sermons
    
    def _sermons_to_documents(self, sermons: List[Dict]) -> List[Document]:
        """Convert sermons to LangChain documents for vector store."""
        documents = []
        
        for sermon in sermons:
            # Combine title and description for embedding
            content = f"{sermon['title']}\n\n{sermon['description']}"
            
            doc = Document(
                page_content=content,
                metadata={
                    'title': sermon['title'],
                    'description': sermon['description'],
                    'channel': sermon['channel'],
                    'message_link': sermon['message_link'],
                    'image_url': sermon.get('image_url'),
                    'date': sermon.get('date'),
                    'theme': sermon.get('theme', '')
                }
            )
            documents.append(doc)
        
        return documents


class MaterialsLoader:
    """Loads documents from materials folder."""
    
    def __init__(self):
        self.db = SermonDatabase(config.DB_PATH)
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200
        )
    
    def load_all_materials(self):
        """Load all files from materials folder."""
        if not os.path.exists(config.MATERIALS_PATH):
            os.makedirs(config.MATERIALS_PATH)
            logger.info(f"Created materials folder: {config.MATERIALS_PATH}")
            return
        
        files = [f for f in os.listdir(config.MATERIALS_PATH) 
                 if f.endswith(('.txt', '.docx', '.pdf'))]
        
        if not files:
            logger.info("No materials to load")
            return
        
        logger.info(f"Loading {len(files)} files from materials folder")
        
        all_documents = []
        for filename in files:
            filepath = os.path.join(config.MATERIALS_PATH, filename)
            
            try:
                # Load file based on extension
                if filename.endswith('.txt'):
                    content = self._load_txt(filepath)
                elif filename.endswith('.docx'):
                    content = self._load_docx(filepath)
                elif filename.endswith('.pdf'):
                    content = self._load_pdf(filepath)
                else:
                    continue
                
                if not content or len(content.strip()) < 50:
                    logger.warning(f"Skipping empty or too short file: {filename}")
                    os.remove(filepath)
                    continue
                
                # Extract metadata from filename or content
                title, link, image = self._parse_filename(filename)
                
                # Create sermon entry
                sermon_data = {
                    'title': title,
                    'description': content[:500],
                    'channel': 'materials',
                    'message_link': link or f"materials/{filename}",
                    'image_url': image,
                    'date': datetime.now().strftime("%Y-%m-%d"),
                    'theme': 'General'
                }
                
                self.db.add_sermon(sermon_data)
                
                # Create documents for vector store
                docs = self.text_splitter.create_documents(
                    [content],
                    metadatas=[{
                        'title': title,
                        'description': content[:500],
                        'channel': 'materials',
                        'message_link': sermon_data['message_link'],
                        'image_url': image,
                        'date': sermon_data['date'],
                        'theme': 'General'
                    }]
                )
                all_documents.extend(docs)
                
                # Delete file after loading
                os.remove(filepath)
                logger.info(f"Loaded and deleted: {filename}")
                
            except Exception as e:
                logger.error(f"Error loading {filename}: {e}")
        
        # Add to vector store
        if all_documents:
            rag_engine.add_documents(all_documents)
            logger.info(f"Added {len(all_documents)} documents to vector store from materials")
    
    def _load_txt(self, filepath: str) -> str:
        """Load text file."""
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(filepath, 'r', encoding='latin-1') as f:
                return f.read()
    
    def _load_docx(self, filepath: str) -> str:
        """Load Word document."""
        try:
            doc = DocxDocument(filepath)
            return '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
        except Exception as e:
            logger.error(f"Error loading DOCX {filepath}: {e}")
            return ""
    
    def _load_pdf(self, filepath: str) -> str:
        """Load PDF file."""
        try:
            reader = PdfReader(filepath)
            text = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
            return '\n'.join(text)
        except Exception as e:
            logger.error(f"Error loading PDF {filepath}: {e}")
            return ""
    
    def _parse_filename(self, filename: str) -> tuple:
        """
        Parse filename for metadata.
        Format: "Title [link] [image.jpg].ext"
        """
        # Remove extension
        name = os.path.splitext(filename)[0]
        
        # Extract link (URL in square brackets)
        link = None
        link_match = re.search(r'\[(https?://[^\]]+)\]', name)
        if link_match:
            link = link_match.group(1)
            name = name.replace(link_match.group(0), '').strip()
        
        # Extract image (filename with image extension in square brackets)
        image = None
        image_match = re.search(r'\[([^\]]+\.(jpg|jpeg|png|gif))\]', name, re.IGNORECASE)
        if image_match:
            image = image_match.group(1)
            name = name.replace(image_match.group(0), '').strip()
        
        # Clean up remaining brackets
        name = re.sub(r'\[|\]', '', name).strip()
        
        # Remaining is title
        title = name if name else "Untitled Sermon"
        
        return title, link, image